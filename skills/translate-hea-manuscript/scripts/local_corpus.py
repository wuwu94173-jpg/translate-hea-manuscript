#!/usr/bin/env python3
"""Index and convert user-authorized local papers without modifying originals."""

from __future__ import annotations

import argparse
import csv
import hashlib
import json
import re
import shutil
import sys
import time
from pathlib import Path
from typing import Any, Iterable

from literature_pipeline import learning_uses, sanitize_filename
from pdf_to_markdown import convert_pdf, yaml_string


SUPPORTED_EXTENSIONS = {".pdf", ".md", ".markdown", ".txt", ".docx"}
DOI_RE = re.compile(r"\b10\.\d{4,9}/[-._;()/:A-Z0-9]+", re.IGNORECASE)
YEAR_RE = re.compile(r"\b(?:19|20)\d{2}\b")


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        while True:
            chunk = handle.read(1024 * 1024)
            if not chunk:
                break
            digest.update(chunk)
    return digest.hexdigest()


def path_is_within(path: Path, parent: Path) -> bool:
    try:
        path.resolve().relative_to(parent.resolve())
        return True
    except ValueError:
        return False


def discover_files(inputs: Iterable[Path], recursive: bool, output_dir: Path) -> list[Path]:
    found: dict[str, Path] = {}
    for supplied in inputs:
        path = supplied.expanduser().resolve()
        if not path.exists():
            print(f"Warning: local input does not exist: {path}", file=sys.stderr)
            continue
        if path.is_file():
            candidates = [path]
        else:
            iterator = path.rglob("*") if recursive else path.glob("*")
            candidates = [candidate for candidate in iterator if candidate.is_file()]
        for candidate in candidates:
            if candidate.suffix.lower() not in SUPPORTED_EXTENSIONS:
                continue
            if path_is_within(candidate, output_dir):
                continue
            found[str(candidate.resolve()).lower()] = candidate.resolve()
    return sorted(found.values(), key=lambda item: str(item).lower())


def read_text_file(path: Path) -> tuple[str, str]:
    raw = path.read_bytes()
    for encoding in ("utf-8-sig", "utf-8", "gb18030", "latin-1"):
        try:
            return raw.decode(encoding), encoding
        except UnicodeDecodeError:
            continue
    raise UnicodeDecodeError("unknown", raw, 0, 1, "no supported encoding")


def first_doi(text: str) -> str:
    match = DOI_RE.search(text or "")
    if not match:
        return ""
    return match.group(0).rstrip(".,;:)]}").lower()


def first_year(text: str) -> str:
    match = YEAR_RE.search(text or "")
    return match.group(0) if match else ""


def markdown_metadata(text: str, fallback_title: str) -> dict[str, Any]:
    frontmatter: dict[str, str] = {}
    if text.startswith("---"):
        parts = text.split("---", 2)
        if len(parts) == 3:
            for line in parts[1].splitlines():
                if ":" in line:
                    key, value = line.split(":", 1)
                    frontmatter[key.strip().lower()] = value.strip().strip('"\'')
    heading = re.search(r"^#\s+(.+)$", text, re.MULTILINE)
    title = frontmatter.get("title") or (heading.group(1).strip() if heading else fallback_title)
    return {
        "title": title,
        "authors": frontmatter.get("authors") or frontmatter.get("author") or "",
        "journal": frontmatter.get("journal", ""),
        "year": frontmatter.get("year") or first_year(text[:5000]),
        "doi": (frontmatter.get("doi") or first_doi(text[:15000])).lower(),
        "text": text,
        "metadata_backend": "markdown/text",
    }


def pdf_metadata(path: Path) -> dict[str, Any]:
    try:
        from pypdf import PdfReader  # type: ignore
    except ImportError as exc:
        raise RuntimeError("pypdf is required for local PDF ingestion") from exc
    reader = PdfReader(str(path))
    metadata = reader.metadata or {}
    sample_parts: list[str] = []
    for page in reader.pages[:3]:
        try:
            sample_parts.append(page.extract_text() or "")
        except Exception:
            continue
    sample = "\n".join(sample_parts)
    raw_title = str(metadata.get("/Title") or "").strip()
    title = raw_title if raw_title and raw_title.lower() not in {"untitled", "unknown"} else path.stem
    created = str(metadata.get("/CreationDate") or "")
    return {
        "title": title,
        "authors": str(metadata.get("/Author") or "").strip(),
        "journal": str(metadata.get("/Subject") or "").strip(),
        "year": first_year(created) or first_year(sample[:5000]),
        "doi": first_doi(sample[:20000]),
        "text": sample,
        "pages": len(reader.pages),
        "metadata_backend": "pypdf",
    }


def docx_metadata(path: Path) -> dict[str, Any]:
    try:
        from docx import Document  # type: ignore
    except ImportError as exc:
        raise RuntimeError("python-docx is required for local DOCX ingestion") from exc
    document = Document(str(path))
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    sample = "\n".join(paragraphs)
    properties = document.core_properties
    title = (properties.title or "").strip() or (paragraphs[0][:300] if paragraphs else path.stem)
    created = properties.created.isoformat() if properties.created else ""
    return {
        "title": title,
        "authors": (properties.author or "").strip(),
        "journal": (properties.subject or "").strip(),
        "year": first_year(created) or first_year(sample[:5000]),
        "doi": first_doi(sample[:20000]),
        "text": sample,
        "metadata_backend": "python-docx",
    }


def extract_metadata(path: Path) -> dict[str, Any]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return pdf_metadata(path)
    if suffix == ".docx":
        return docx_metadata(path)
    text, encoding = read_text_file(path)
    metadata = markdown_metadata(text, path.stem)
    metadata["source_encoding"] = encoding
    return metadata


def render_imported_text(record: dict[str, Any], text: str, source_type: str) -> str:
    frontmatter = [
        "---",
        f"title: {yaml_string(record.get('title'))}",
        f"authors: {yaml_string(record.get('authors'))}",
        f"journal: {yaml_string(record.get('journal'))}",
        f"year: {yaml_string(record.get('year'))}",
        f"doi: {yaml_string(record.get('doi'))}",
        f"source_origin: {yaml_string('local')}",
        f"source_local_path: {yaml_string(record.get('source_path'))}",
        f"source_sha256: {yaml_string(record.get('sha256'))}",
        f"source_type: {yaml_string(source_type)}",
        f"main_translation_use: {yaml_string(', '.join(record.get('main_translation_use') or []))}",
        "---",
        "",
        "> Imported from a user-authorized local file. The original file was not modified.",
        "",
    ]
    return "\n".join(frontmatter) + text.strip() + "\n"


def docx_to_markdown(path: Path) -> str:
    from docx import Document  # type: ignore

    document = Document(str(path))
    lines: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = (paragraph.style.name or "").lower() if paragraph.style else ""
        heading = re.match(r"heading\s+(\d+)", style_name)
        if heading:
            level = min(max(int(heading.group(1)), 1), 6)
            lines.extend([f"{'#' * level} {text}", ""])
        else:
            lines.extend([text, ""])
    for table_number, table in enumerate(document.tables, 1):
        rows = [[cell.text.strip().replace("|", "\\|") for cell in row.cells] for row in table.rows]
        if not rows:
            continue
        width = max(len(row) for row in rows)
        rows = [row + [""] * (width - len(row)) for row in rows]
        lines.extend([f"### Imported table {table_number}", ""])
        lines.append("| " + " | ".join(rows[0]) + " |")
        lines.append("| " + " | ".join(["---"] * width) + " |")
        for row in rows[1:]:
            lines.append("| " + " | ".join(row) + " |")
        lines.append("")
    return "\n".join(lines).strip()


def copy_original(path: Path, destination_dir: Path, digest: str, overwrite: bool) -> str:
    name = f"{sanitize_filename(path.stem, 80)}_{digest[:8]}{path.suffix.lower()}"
    destination = destination_dir / name
    if not destination.exists() or overwrite:
        destination_dir.mkdir(parents=True, exist_ok=True)
        shutil.copy2(path, destination)
    return name


def import_record(
    path: Path,
    index: int,
    markdown_dir: Path,
    originals_dir: Path,
    copy_originals: bool,
    overwrite: bool,
) -> dict[str, Any]:
    digest = sha256_file(path)
    record: dict[str, Any] = {
        "source_path": str(path),
        "source_filename": path.name,
        "source_type": path.suffix.lower().lstrip("."),
        "size_bytes": path.stat().st_size,
        "modified_time": time.strftime("%Y-%m-%dT%H:%M:%S", time.localtime(path.stat().st_mtime)),
        "sha256": digest,
        "title": path.stem,
        "authors": "",
        "journal": "",
        "year": "",
        "doi": "",
        "metadata_status": "pending",
        "metadata_error": "",
        "duplicate_of": "",
        "copied_filename": "",
        "conversion_status": "pending",
        "conversion_error": "",
        "markdown_filename": "",
        "main_translation_use": [],
    }
    try:
        metadata = extract_metadata(path)
        record.update({key: value for key, value in metadata.items() if key != "text"})
        record["metadata_status"] = "extracted"
        record["main_translation_use"] = learning_uses(
            f"{record.get('title', '')} {metadata.get('text', '')[:30000]}"
        )
    except Exception as exc:
        metadata = {"text": ""}
        record["metadata_status"] = "failed"
        record["metadata_error"] = f"{type(exc).__name__}: {exc}"

    if copy_originals:
        try:
            record["copied_filename"] = copy_original(path, originals_dir, digest, overwrite)
        except Exception as exc:
            record["copy_error"] = f"{type(exc).__name__}: {exc}"

    stem = f"local_{index:02d}_{sanitize_filename(str(record.get('title') or path.stem), 70)}_{digest[:8]}"
    markdown_path = markdown_dir / f"{stem}.md"
    try:
        if path.suffix.lower() == ".pdf":
            conversion_metadata = {
                "title": record.get("title"),
                "authors": record.get("authors"),
                "journal": record.get("journal"),
                "year": record.get("year"),
                "doi": record.get("doi"),
                "source_origin": "local",
                "source_local_path": record.get("source_path"),
                "source_sha256": record.get("sha256"),
                "main_translation_use": ", ".join(record.get("main_translation_use") or []),
                "language_quality": "pending agent review",
            }
            report = convert_pdf(path, markdown_path, conversion_metadata, overwrite)
            record["conversion_report"] = report
            record["conversion_status"] = report.get("status", "converted")
        else:
            if markdown_path.exists() and not overwrite:
                record["conversion_status"] = "exists"
            else:
                if path.suffix.lower() == ".docx":
                    body = docx_to_markdown(path)
                else:
                    body = metadata.get("text") or read_text_file(path)[0]
                markdown_path.parent.mkdir(parents=True, exist_ok=True)
                markdown_path.write_text(
                    render_imported_text(record, body, path.suffix.lower().lstrip(".")),
                    encoding="utf-8",
                )
                record["conversion_status"] = "converted"
        record["markdown_filename"] = markdown_path.name
    except Exception as exc:
        record["conversion_status"] = "failed"
        record["conversion_error"] = f"{type(exc).__name__}: {exc}"
    return record


def mark_duplicates(records: list[dict[str, Any]]) -> None:
    hashes: dict[str, dict[str, Any]] = {}
    dois: dict[str, dict[str, Any]] = {}
    for record in records:
        duplicate: dict[str, Any] | None = None
        digest = record.get("sha256")
        doi = str(record.get("doi") or "").lower()
        if digest in hashes:
            duplicate = hashes[digest]
        elif doi and doi in dois:
            duplicate = dois[doi]
        if duplicate:
            record["duplicate_of"] = duplicate.get("source_path", "")
            record["duplicate_reason"] = "identical_sha256" if digest == duplicate.get("sha256") else "same_doi"
        else:
            hashes[digest] = record
            if doi:
                dois[doi] = record


def markdown_escape(value: Any) -> str:
    return re.sub(r"\s+", " ", "" if value is None else str(value)).strip().replace("|", "\\|")


def write_indexes(output_dir: Path, records: list[dict[str, Any]], inputs: list[Path]) -> None:
    reference_dir = output_dir / "02_reference_papers"
    reference_dir.mkdir(parents=True, exist_ok=True)
    manifest = {
        "generated_at_utc": time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime()),
        "source_inputs": [str(path.resolve()) for path in inputs],
        "source_files_modified": False,
        "papers": records,
    }
    (reference_dir / "local_corpus_manifest.json").write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    fields = [
        "source_path", "source_filename", "source_type", "size_bytes", "sha256",
        "title", "authors", "journal", "year", "doi", "metadata_status",
        "metadata_error", "duplicate_of", "duplicate_reason", "copied_filename",
        "conversion_status", "conversion_error", "markdown_filename",
        "main_translation_use_text",
    ]
    with (reference_dir / "local_literature_index.csv").open(
        "w", encoding="utf-8-sig", newline=""
    ) as handle:
        writer = csv.DictWriter(handle, fieldnames=fields)
        writer.writeheader()
        for record in records:
            row = dict(record)
            row["main_translation_use_text"] = "; ".join(record.get("main_translation_use") or [])
            writer.writerow({field: row.get(field, "") for field in fields})

    lines = [
        "# Local literature index", "",
        "> Original local files were read only. Duplicate records are retained in the index but should be learned once.",
        "", "| File | Title | Year | DOI | Type | Duplicate | Conversion | Translation use |",
        "|---|---|---:|---|---|---|---|---|",
    ]
    for record in records:
        lines.append(
            "| {file} | {title} | {year} | {doi} | {type} | {duplicate} | {status} | {uses} |".format(
                file=markdown_escape(record.get("source_filename")),
                title=markdown_escape(record.get("title")),
                year=markdown_escape(record.get("year")),
                doi=markdown_escape(record.get("doi")),
                type=markdown_escape(record.get("source_type")),
                duplicate="yes" if record.get("duplicate_of") else "no",
                status=markdown_escape(record.get("conversion_status")),
                uses=markdown_escape(", ".join(record.get("main_translation_use") or [])),
            )
        )
    (reference_dir / "local_literature_index.md").write_text(
        "\n".join(lines) + "\n", encoding="utf-8"
    )

    queue = [
        "# Local literature learning queue", "",
        "Process each non-duplicate converted paper once. Preserve its local source hash and use DOI/page anchors when available.",
        "",
    ]
    for record in records:
        if record.get("duplicate_of") or not record.get("markdown_filename"):
            continue
        queue.extend([
            f"## {record.get('title')}", "",
            f"- Markdown: `markdown/{record.get('markdown_filename')}`",
            f"- Original: `{record.get('source_path')}`",
            f"- SHA-256: `{record.get('sha256')}`",
            f"- DOI: `{record.get('doi') or 'not detected'}`",
            f"- Translation use: {', '.join(record.get('main_translation_use') or [])}",
            "- Validate title/DOI metadata, conversion quality, relevance, and page anchors before learning.",
            "",
        ])
    (reference_dir / "local_learning_queue.md").write_text(
        "\n".join(queue), encoding="utf-8"
    )


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Read-only ingestion of local PDF, Markdown, text, and DOCX papers."
    )
    parser.add_argument("inputs", nargs="+", type=Path, help="Local paper files or directories")
    parser.add_argument("--output-dir", required=True, type=Path, help="HEA translation project directory")
    parser.add_argument("--recursive", action="store_true", help="Scan input directories recursively")
    parser.add_argument("--copy-originals", action="store_true", help="Copy originals into the project; never moves them")
    parser.add_argument("--overwrite", action="store_true", help="Replace generated copies and Markdown")
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    output_dir = args.output_dir.resolve()
    files = discover_files(args.inputs, args.recursive, output_dir)
    if not files:
        print("No supported local papers found.", file=sys.stderr)
        write_indexes(output_dir, [], args.inputs)
        return 1
    markdown_dir = output_dir / "02_reference_papers" / "markdown"
    originals_dir = output_dir / "02_reference_papers" / "local_originals"
    markdown_dir.mkdir(parents=True, exist_ok=True)
    records = [
        import_record(
            path, index, markdown_dir, originals_dir,
            args.copy_originals, args.overwrite,
        )
        for index, path in enumerate(files, 1)
    ]
    mark_duplicates(records)
    write_indexes(output_dir, records, args.inputs)
    converted = sum(record.get("conversion_status") in {"converted", "exists"} for record in records)
    duplicates = sum(bool(record.get("duplicate_of")) for record in records)
    print(f"Local papers: {len(records)}; converted/indexed: {converted}; duplicates: {duplicates}")
    print(output_dir / "02_reference_papers" / "local_literature_index.md")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
